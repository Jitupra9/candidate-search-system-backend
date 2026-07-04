"""
DocumentLoader
==============
Supports:
  .pdf   → plain text PDF  (PyMuPDF)
  .pdf   → tabular PDF     (pdfplumber — detects tables automatically)
  .docx  → Word document   (python-docx)
  .doc   → legacy Word     (textract fallback)
  .csv   → CSV file        (pandas)
  .xlsx  → Excel file      (pandas)
  .txt   → plain text      (built-in)

All loaders return List[Document] with metadata:
  { source, file_type, page (if applicable), row_index (CSV/Excel) }
"""
import logging
import re
from pathlib import Path
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ── Text cleanup ────────────────────────────────────────────────────────────

# Matches PDF font glyph artifacts like (cid:127), (cid:9) etc. that PyMuPDF /
# pdfplumber emit when a custom bullet/symbol glyph has no Unicode mapping.
_CID_PATTERN = re.compile(r"\(cid:\d+\)")

# A line is treated as a "section heading" if it is short, has no lowercase
# letters at all (i.e. ALL CAPS), and doesn't look like a sentence.
_HEADING_PATTERN = re.compile(r"^[A-Z0-9][A-Z0-9 &/\-,]{2,40}$")


def _normalize_text(text: str) -> str:
    """Strip PDF glyph artifacts and normalize bullet markers."""
    if not text:
        return text
    # Replace cid-encoded bullet glyphs with a plain hyphen bullet.
    text = _CID_PATTERN.sub("-", text)
    # Collapse the now-common "-  Text" (extra spaces left behind) into "- Text"
    text = re.sub(r"^-\s+", "- ", text, flags=re.MULTILINE)
    # Collapse 3+ blank lines into a single blank line
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _is_heading_line(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 40:
        return False
    return bool(_HEADING_PATTERN.match(line))


# ── PDF ───────────────────────────────────────────────────────────────────────

def _clean_cell(cell) -> str:
    if cell is None:
        return ""
    cleaned = _CID_PATTERN.sub("-", str(cell))
    return re.sub(r"\s+", " ", cleaned).strip()


def _is_numeric_like(s: str) -> bool:
    s = s.strip().replace(",", "").replace("%", "").replace("$", "")
    if not s:
        return False
    return bool(re.match(r"^-?\d+(\.\d+)?$", s))


def _looks_like_header(header_row: list[str], data_rows: list[list[str]]) -> bool:
    if not header_row or not data_rows:
        return False
    if not all(not cell or len(cell) <= 30 for cell in header_row):
        return False
    if any(_is_numeric_like(cell) for cell in header_row):
        return False

    num_cols = len(header_row)
    for col_idx in range(num_cols):
        col_values = [row[col_idx] for row in data_rows if col_idx < len(row)]
        if not col_values:
            continue
        numeric_ratio = sum(_is_numeric_like(v) for v in col_values) / len(col_values)
        if numeric_ratio >= 0.5:
            return True

    return False


def _table_to_text(table_obj) -> str:
    rows = table_obj.extract()
    rows = [[_clean_cell(c) for c in r] for r in rows]
    if not rows:
        return ""

    header_candidate, data_rows = rows[0], rows[1:]

    if _looks_like_header(header_candidate, data_rows):
        header = [h if h else f"col{i}" for i, h in enumerate(header_candidate)]
        lines = []
        for row in data_rows:
            row_text = ", ".join(f"{h}: {v}" for h, v in zip(header, row) if v)
            if row_text:
                lines.append(row_text)
        return "\n".join(lines)

    return "\n".join(" | ".join(r) for r in rows)


def _extract_page_in_order(page) -> tuple[str, bool]:
    table_objs = sorted(page.find_tables(), key=lambda t: t.bbox[1])

    if not table_objs:
        text = (page.extract_text() or "").strip()
        return _normalize_text(text), False

    parts: list[str] = []
    cursor_top = 0.0
    page_bottom = page.height

    for t in table_objs:
        _, top, _, bottom = t.bbox
        top = max(0.0, min(top, page_bottom))

        if top > cursor_top:
            strip = page.crop((0, cursor_top, page.width, top))
            strip_text = (strip.extract_text() or "").strip()
            if strip_text:
                parts.append(_normalize_text(strip_text))

        table_text = _table_to_text(t)
        if table_text:
            parts.append(_normalize_text(table_text))

        cursor_top = max(cursor_top, min(bottom, page_bottom))

    if cursor_top < page_bottom:
        strip = page.crop((0, cursor_top, page.width, page_bottom))
        strip_text = (strip.extract_text() or "").strip()
        if strip_text:
            parts.append(_normalize_text(strip_text))

    return "\n\n".join(p for p in parts if p), True


def _load_pdf(path: str) -> list[Document]:
    import pdfplumber
    from langchain_community.document_loaders import PyMuPDFLoader

    docs: list[Document] = []

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text, has_table = _extract_page_in_order(page)

            if page_text.strip():
                docs.append(Document(
                    page_content=page_text.strip(),
                    metadata={
                        "source": path,
                        "file_type": "pdf",
                        "page": page_num + 1,
                        "has_table": has_table,
                    },
                ))

    if not docs:
        logger.warning("pdfplumber got no text, falling back to PyMuPDF: %s", path)
        loader = PyMuPDFLoader(path)
        raw = loader.load()
        docs = [
            Document(
                page_content=_normalize_text(d.page_content),
                metadata={**d.metadata, "file_type": "pdf", "has_table": False},
            )
            for d in raw
        ]

    return docs


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _load_docx(path: str) -> list[Document]:
    """
    Word document loader.
    - Extracts paragraphs as text
    - Extracts tables as pipe-separated rows
    - Groups into logical sections (heading → content)
    """
    from docx import Document as DocxDocument

    docx = DocxDocument(path)
    sections: list[str] = []
    current_section = ""

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # New heading → save previous section, start new one
        if para.style.name.startswith("Heading"):
            if current_section:
                sections.append(current_section.strip())
            current_section = f"## {text}\n"
        else:
            current_section += text + "\n"

    if current_section:
        sections.append(current_section.strip())

    # Extract tables
    for table in docx.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            sections.append("\n".join(rows))

    return [
        Document(
            page_content=_normalize_text(section),
            metadata={"source": path, "file_type": "docx", "section_index": i},
        )
        for i, section in enumerate(sections)
        if section.strip()
    ]


# ── CSV ───────────────────────────────────────────────────────────────────────

def _detect_encoding(file_path: str) -> str:
    """
    Detect file encoding by trying common encodings.
    Returns the first successful encoding.
    """
    encodings = [
        'utf-8',
        'utf-8-sig',  # UTF-8 with BOM
        'windows-1252',  # Common Windows encoding
        'latin-1',  # ISO-8859-1
        'cp1252',  # Another Windows encoding
        'iso-8859-1',
        'mac_roman',  # Mac encoding
        'ascii',
    ]

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                f.read()
            return encoding
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Fallback to latin-1 which never fails (maps all bytes to characters)
    return 'latin-1'


def _load_csv(path: str) -> list[Document]:
    """
    CSV loader with robust encoding detection.
    Each row becomes a Document with column context.
    """
    import pandas as pd

    # Detect encoding first
    encoding = _detect_encoding(path)
    logger.debug(f"Detected encoding for {path}: {encoding}")

    try:
        # Try with detected encoding
        df = pd.read_csv(path, encoding=encoding)
    except UnicodeDecodeError:
        # If detection fails, try reading as binary and handling errors
        try:
            with open(path, 'rb') as f:
                content = f.read()
                # Try to decode with error replacement
                text = content.decode('utf-8', errors='replace')
                from io import StringIO
                df = pd.read_csv(StringIO(text))
        except Exception as e:
            logger.error(f"Failed to read CSV with fallback: {e}")
            # Final fallback: read with pandas using 'latin-1'
            df = pd.read_csv(path, encoding='latin-1')

    df = df.fillna("")
    docs: list[Document] = []

    # Try to extract candidate name if column exists
    name_columns = ['name', 'Name', 'full_name', 'Full Name', 'candidate_name', 'candidate']
    name_col = None
    for col in name_columns:
        if col in df.columns:
            name_col = col
            break

    for idx, row in df.iterrows():
        row_text = ", ".join(
            f"{col}: {str(val).strip()}"
            for col, val in row.items()
            if str(val).strip()
        )
        if row_text:
            # Build metadata
            metadata = {
                "source": path,
                "file_type": "csv",
                "row_index": int(idx),
                "columns": list(df.columns),
            }

            # Add candidate name if available
            if name_col:
                metadata["candidate_name"] = str(row.get(name_col, f"Row_{idx}"))

            docs.append(Document(
                page_content=row_text,
                metadata=metadata,
            ))

    logger.info(f"Loaded {len(docs)} rows from CSV: {Path(path).name}")
    return docs


# ── XLSX ──────────────────────────────────────────────────────────────────────

def _load_xlsx(path: str) -> list[Document]:
    """
    Excel loader — processes all sheets.
    Each row becomes a Document with sheet name in metadata.
    """
    import pandas as pd

    xl = pd.ExcelFile(path)
    docs: list[Document] = []

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name).fillna("")
        for idx, row in df.iterrows():
            row_text = ", ".join(
                f"{col}: {str(val).strip()}"
                for col, val in row.items()
                if str(val).strip()
            )
            if row_text:
                docs.append(Document(
                    page_content=row_text,
                    metadata={
                        "source": path,
                        "file_type": "xlsx",
                        "sheet": sheet_name,
                        "row_index": int(idx),
                    },
                ))

    return docs


# ── TXT ───────────────────────────────────────────────────────────────────────

def _load_txt(path: str) -> list[Document]:
    """Plain text — split on double newlines into paragraphs."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    return [
        Document(
            page_content=_normalize_text(content),
            metadata={"source": path, "file_type": "txt"},
        )
    ]


# ── Router ────────────────────────────────────────────────────────────────────

_LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".doc":  _load_docx,   # python-docx handles .doc too in most cases
    ".csv":  _load_csv,
    ".xlsx": _load_xlsx,
    ".xls":  _load_xlsx,
    ".txt":  _load_txt,
}


def document_Loader(document_path: str) -> list[Document]:
    """
    Main entry point.
    Auto-detects file type and routes to correct loader.
    Returns List[Document] with normalized metadata.

    Raises:
        ValueError  — unsupported file type
        RuntimeError — loader failure with cause
    """
    path = Path(document_path)
    ext = path.suffix.lower()

    loader_fn = _LOADERS.get(ext)
    if not loader_fn:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {list(_LOADERS.keys())}"
        )

    try:
        docs = loader_fn(document_path)
        logger.info("Loaded %d documents from %s (%s)", len(docs), path.name, ext)
        return docs
    except Exception as e:
        logger.error("Failed to load %s: %s", path.name, e)
        raise RuntimeError(f"Failed to load '{path.name}': {e}") from e